"""
Resume Parser Module
Parses resume from various formats (PDF, DOCX, TXT)
"""

import os
import re
import json
from typing import Dict, Optional, Union
from io import BytesIO


class ResumeParser:
    """Parse resume from different file formats."""

    def __init__(self, llm_client: Optional[any] = None):
        self.supported_formats = [".pdf", ".docx", ".txt"]
        self.llm_client = llm_client

    def parse(self, file_input) -> Dict[str, any]:
        """Parse a resume file and extract structured information.

        Args:
            file_input: Path to the resume file or Streamlit UploadedFile

        Returns:
            Dictionary containing parsed resume data
        """
        # Handle Streamlit UploadedFile
        if hasattr(file_input, "read"):
            # It's a Streamlit UploadedFile
            file_name = getattr(file_input, "name", "resume.pdf")
            ext = os.path.splitext(file_name)[1].lower()

            if ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {ext}")

            # Read content
            content = file_input.read()

            if ext == ".pdf":
                result = self._parse_pdf_bytes(content)
            elif ext == ".docx":
                result = self._parse_docx_bytes(content)
            elif ext == ".txt":
                result = self._parse_txt_bytes(content)
        else:
            # Handle file path (original behavior)
            file_path = file_input
            ext = os.path.splitext(file_path)[1].lower()

            if ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {ext}")

            if ext == ".pdf":
                result = self._parse_pdf(file_input)
            elif ext == ".docx":
                result = self._parse_docx(file_input)
            elif ext == ".txt":
                result = self._parse_txt(file_input)

        # If LLM client is available, use it for better extraction
        if self.llm_client and result and not result.get("error"):
            return self._ai_extract_sections(result.get("raw_text", ""))

        return result

    def _ai_extract_sections(self, raw_text: str) -> Dict[str, any]:
        """Use AI to extract structured sections from resume text."""
        prompt = f"""Extract structured information from the following resume text.
Return a STRICT JSON object with these keys:
- "name": Full name
- "email": Email address
- "phone": Phone number
- "linkedin": LinkedIn URL
- "summary": Professional summary
- "experience": List of jobs, each with "title", "company", "dates", and "description" (as a string with bullet points)
- "education": List of degrees/certificates
- "skills": List of skills

RESUME TEXT:
{raw_text}
"""
        try:
            # Use chat interface of UniversalLLMClient
            response = self.llm_client.chat([{"role": "user", "content": prompt}])

            # Clean response if LLM added markdown backticks
            if "```json" in response:
                response = response.split("```json")[1].split("```")[0].strip()
            elif "```" in response:
                response = response.split("```")[1].split("```")[0].strip()

            structured_data = json.loads(response)
            structured_data["raw_text"] = raw_text
            return structured_data
        except Exception as e:
            logger.error(f"AI Extraction Error: {e}")
            # Fallback to regex extraction if AI fails
            return self._extract_sections(raw_text)

    def _standardize_bullets(self, text: str) -> str:
        """Standardize various bullet characters to a single format."""
        if not text:
            return ""
        # Standardize various bullet characters to '•'
        bullet_chars = "●○▪■◦·‣⁃"
        for char in bullet_chars:
            text = text.replace(char, "•")
        
        # Handle hyphens and asterisks used as bullets at start of line
        # This regex looks for -, * or • preceded by whitespace at the start of a line
        text = re.sub(r"^[ \t]*[•\-*][ \t]+", "• ", text, flags=re.MULTILINE)
        return text

    def _parse_pdf_bytes(self, content: bytes) -> Dict[str, any]:
        """Parse PDF from bytes."""
        text = ""

        try:
            import pdfplumber

            with pdfplumber.open(BytesIO(content)) as pdf:
                for page in pdf.pages:
                    # Use layout=True to preserve visual structure
                    text += page.extract_text(layout=True) or ""
                    text += "\n"
        except ImportError:
            try:
                from pypdf import PdfReader

                reader = PdfReader(BytesIO(content))
                for page in reader.pages:
                    text += page.extract_text() or ""
                    text += "\n"
            except Exception as e:
                return {"error": f"Failed to parse PDF: {str(e)}", "raw_text": ""}
        except Exception as e:
            return {"error": f"Failed to parse PDF: {str(e)}", "raw_text": ""}

        text = self._standardize_bullets(text)
        return self._extract_sections(text)

    def _parse_docx_bytes(self, content: bytes) -> Dict[str, any]:
        """Parse DOCX from bytes."""
        try:
            from docx import Document

            doc = Document(BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs])
            text = self._standardize_bullets(text)
            return self._extract_sections(text)
        except Exception as e:
            return {"error": f"Failed to parse DOCX: {str(e)}", "raw_text": ""}

    def _parse_txt_bytes(self, content: bytes) -> Dict[str, any]:
        """Parse TXT from bytes."""
        try:
            text = content.decode("utf-8")
            text = self._standardize_bullets(text)
            return self._extract_sections(text)
        except Exception as e:
            return {"error": f"Failed to parse TXT: {str(e)}", "raw_text": ""}

    def _parse_pdf(self, file_path: str) -> Dict[str, any]:
        """Parse PDF resume using pdfplumber or pypdf."""
        text = ""

        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    # Use layout=True to preserve visual structure
                    text += page.extract_text(layout=True) or ""
                    text += "\n"
        except ImportError:
            try:
                from pypdf import PdfReader

                reader = PdfReader(file_path)
                for page in reader.pages:
                    text += page.extract_text() or ""
                    text += "\n"
            except Exception as e:
                return {"error": f"Failed to parse PDF: {str(e)}", "raw_text": ""}
        except Exception as e:
            return {"error": f"Failed to parse PDF: {str(e)}", "raw_text": ""}

        text = self._standardize_bullets(text)
        return self._extract_sections(text)

    def _parse_docx(self, file_path: str) -> Dict[str, any]:
        """Parse DOCX resume using python-docx."""
        try:
            from docx import Document

            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            text = self._standardize_bullets(text)
            return self._extract_sections(text)
        except Exception as e:
            return {"error": f"Failed to parse DOCX: {str(e)}", "raw_text": ""}

    def _parse_txt(self, file_path: str) -> Dict[str, any]:
        """Parse TXT resume."""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            text = self._standardize_bullets(text)
            return self._extract_sections(text)
        except Exception as e:
            return {"error": f"Failed to parse TXT: {str(e)}", "raw_text": ""}

    def _extract_sections(self, text: str) -> Dict[str, any]:
        """Extract structured sections from resume text."""
        # Preserve original text with newlines for ATS scoring
        raw_text = text
        
        # Create a collapsed version for easier regex matching
        lines = text.split("\n")
        collapsed_text = " ".join([line.strip() for line in lines if line.strip()])

        sections = {
            "name": self._extract_name(lines),
            "email": self._extract_email(collapsed_text),
            "phone": self._extract_phone(collapsed_text),
            "linkedin": self._extract_linkedin(collapsed_text),
            "summary": self._extract_summary(collapsed_text),
            "experience": self._extract_experience(text),
            "education": self._extract_education(collapsed_text),
            "skills": self._extract_skills(text),
            "raw_text": raw_text,
        }

        return sections

    def _extract_name(self, lines: list) -> Optional[str]:
        """Extract name from resume."""
        for line in lines:
            line = line.strip()
            if not line or len(line) > 50:
                continue
            if "@" in line or any(p in line.lower() for p in ["linkedin.com", "phone:", "email:"]):
                continue
            return line
        return None

    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address."""
        pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        match = re.search(pattern, text)
        return match.group(0) if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number."""
        pattern = r"\b[\+]?[(]?[0-9]{1,3}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}\b"
        match = re.search(pattern, text)
        return match.group(0) if match else None

    def _extract_linkedin(self, text: str) -> Optional[str]:
        """Extract LinkedIn URL."""
        pattern = r"linkedin\.com/in/[a-zA-Z0-9-]+"
        match = re.search(pattern, text, re.IGNORECASE)
        return match.group(0) if match else None

    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract professional summary."""
        patterns = [
            r"professional summary[:\s]+(.*?)(?=\b(?:experience|education|skills|projects)\b|$)",
            r"summary[:\s]+(.*?)(?=\b(?:experience|education|skills|projects)\b|$)",
            r"about[:\s]+(.*?)(?=\b(?:experience|education|skills|projects)\b|$)",
        ]
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                return match.group(1).strip()
        return None

    def _extract_experience(self, text: str) -> list:
        """Extract work experience."""
        experience = []
        
        # Identify the experience section
        exp_section_pattern = r"(?:experience|work history|professional background|employment)[:\s]+(.*?)(?=\b(?:education|skills|projects|certifications|awards)\b|$)"
        match = re.search(exp_section_pattern, text, re.IGNORECASE | re.DOTALL)
        
        content = match.group(1) if match else text
        lines = content.split("\n")
        
        current_job = None
        date_pattern = r"\b(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|present|[0-9]{4})\b"
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            has_date = len(re.findall(date_pattern, line, re.IGNORECASE)) >= 1
            
            if not line.startswith("•") and (has_date or len(line) < 100):
                if current_job:
                    # Convert description list to string for compatibility
                    current_job["description"] = "\n".join(current_job["description"])
                    experience.append(current_job)
                
                current_job = {
                    "title": line,
                    "company": "",
                    "dates": "",
                    "description": []
                }
                date_match = re.search(r"((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|present|[0-9]{2}/[0-9]{4}|[0-9]{4})\s*(?:-|–|to)\s*(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|present|[0-9]{2}/[0-9]{4}|[0-9]{4}))", line, re.IGNORECASE)
                if date_match:
                    current_job["dates"] = date_match.group(1)
                    current_job["title"] = line.replace(date_match.group(1), "").strip().strip("|").strip()
            
            elif line.startswith("•") and current_job:
                current_job["description"].append(line)
            
            elif current_job and not current_job["company"] and not line.startswith("•"):
                current_job["company"] = line

        if current_job:
            current_job["description"] = "\n".join(current_job["description"])
            experience.append(current_job)
            
        if not experience:
             job_patterns = [r"Product Manager", r"Engineer", r"Analyst", r"Director", r"Lead", r"Consultant"]
             for i, line in enumerate(lines):
                 for pattern in job_patterns:
                     if re.search(pattern, line, re.IGNORECASE) and not line.startswith("•"):
                         experience.append({"title": line, "company": "", "dates": "", "description": ""})
                         break

        return experience

    def _extract_education(self, text: str) -> list:
        """Extract education."""
        education = []
        patterns = [
            r"Bachelor[s]?\s*(?:of)?\s*([A-Za-z\s]+)",
            r"Master[s]?\s*(?:of)?\s*([A-Za-z\s]+)",
            r"B\.\w+",
            r"M\.\w+",
            r"Ph\.?D\.?",
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if match:
                    education.append(match.strip())
        return education

    def _extract_skills(self, text: str) -> list:
        """Extract skills."""
        skills = []
        skills_pattern = r"(?:skills|technical skills|competencies|technologies)[:\s]+(.*?)(?=\b(?:experience|education|projects|certifications)\b|$)"
        match = re.search(skills_pattern, text, re.IGNORECASE | re.DOTALL)
        
        if match:
            skills_text = match.group(1)
            delimiters = [",", "|", "•", "\n"]
            for d in delimiters:
                if d in skills_text:
                    parts = [p.strip() for p in skills_text.split(d) if p.strip()]
                    skills.extend(parts)
                    break
            else:
                skills.append(skills_text.strip())

        return list(set([s.strip() for s in skills if s.strip() and len(s) < 50]))


def parse_resume(file_path: str) -> Dict[str, any]:
    """Convenience function to parse a resume."""
    parser = ResumeParser()
    return parser.parse(file_path)
